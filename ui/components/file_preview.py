"""Secure in-app preview dialog for authorized company-form files.

The caller must obtain file bytes through the service layer first. This
component only renders those already-authorized bytes and never reads an
arbitrary client path.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document


_MAX_PREVIEW_ROWS = 250
_MAX_PREVIEW_COLUMNS = 40


def _read_text(data: bytes) -> str:
    """Decode text safely using common encodings."""

    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _render_pdf(data: bytes) -> None:
    """Render PDF bytes with Streamlit's native in-app PDF viewer."""

    # The old PDF-in-an-iframe approach could render a blank white area in
    # Chromium browsers. ``st.pdf`` uses Streamlit's supported PDF component
    # and accepts authorized raw bytes directly.
    st.pdf(data, height=650)


def _render_docx(data: bytes) -> None:
    """Render DOCX paragraphs and tables in a readable modal layout."""

    document = Document(BytesIO(data))
    has_content = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            has_content = True
            st.markdown(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        has_content = True
        st.caption(f"Table {table_index}")
        st.dataframe(
            pd.DataFrame(rows),
            hide_index=True,
            use_container_width=True,
            height=min(360, 38 + (len(rows) * 35)),
        )

    if not has_content:
        st.info("The DOCX file does not contain previewable text or tables.")


def _render_spreadsheet(data: bytes, filename: str) -> None:
    """Render a selectable sheet from a modern Excel workbook."""

    workbook = pd.ExcelFile(BytesIO(data), engine="openpyxl")
    sheet_name = st.selectbox(
        "Worksheet",
        options=workbook.sheet_names,
        key=f"company_file_preview_sheet_{abs(hash(filename))}",
    )
    frame = workbook.parse(sheet_name=sheet_name, header=None)
    clipped = frame.iloc[:_MAX_PREVIEW_ROWS, :_MAX_PREVIEW_COLUMNS]
    st.dataframe(
        clipped,
        hide_index=True,
        use_container_width=True,
        height=560,
    )
    if frame.shape != clipped.shape:
        st.caption(
            f"Preview is limited to {_MAX_PREVIEW_ROWS} rows and "
            f"{_MAX_PREVIEW_COLUMNS} columns. Download the file for the "
            "complete workbook."
        )


def _render_csv(data: bytes) -> None:
    """Render a CSV table with a safe preview size."""

    frame = pd.read_csv(BytesIO(data))
    clipped = frame.iloc[:_MAX_PREVIEW_ROWS, :_MAX_PREVIEW_COLUMNS]
    st.dataframe(
        clipped,
        hide_index=True,
        use_container_width=True,
        height=560,
    )
    if frame.shape != clipped.shape:
        st.caption(
            f"Preview is limited to {_MAX_PREVIEW_ROWS} rows and "
            f"{_MAX_PREVIEW_COLUMNS} columns."
        )


def _render_image(data: bytes) -> None:
    """Render an uploaded image submission."""

    st.image(data, use_container_width=True)


def _render_preview(*, filename: str, mime_type: str, data: bytes) -> None:
    """Dispatch one authorized file to the matching preview renderer."""

    extension = Path(filename).suffix.lower()

    if extension == ".pdf" or mime_type == "application/pdf":
        _render_pdf(data)
        return
    if extension == ".docx":
        _render_docx(data)
        return
    if extension == ".xlsx":
        _render_spreadsheet(data, filename)
        return
    if extension == ".csv":
        _render_csv(data)
        return
    if extension == ".txt" or mime_type.startswith("text/"):
        st.text_area(
            "Text Preview",
            value=_read_text(data),
            height=560,
            disabled=True,
            label_visibility="collapsed",
        )
        return
    if extension in {".png", ".jpg", ".jpeg"} or mime_type.startswith("image/"):
        _render_image(data)
        return

    if extension in {".doc", ".xls"}:
        st.info(
            "Legacy Word (.doc) and Excel (.xls) files cannot be rendered "
            "reliably in the browser. Download the file to open it in the "
            "appropriate desktop application."
        )
        return

    st.info(
        "A browser preview is not available for this file type. Use the "
        "download button below."
    )


@st.dialog("File Preview", width="large", dismissible=False)
def render_file_preview_dialog(
    *,
    filename: str,
    mime_type: str,
    data: bytes,
    preview_state_key: str,
    table_version_key: str,
) -> None:
    """Show one authorized file in a large modal with download fallback."""

    st.subheader(filename)
    st.caption("Preview only. The stored company file is not modified.")

    try:
        _render_preview(filename=filename, mime_type=mime_type, data=data)
    except Exception as error:  # A malformed document must not break the page.
        st.warning(
            "This file could not be rendered in the browser preview. "
            "Download it to open the complete file."
        )
        st.caption(f"Preview detail: {error}")

    st.download_button(
        "Download File",
        data=data,
        file_name=filename,
        mime=mime_type,
        type="primary",
        use_container_width=True,
        key=f"preview_download_{preview_state_key}_{abs(hash(filename))}",
    )

    if st.button(
        "Close Preview",
        use_container_width=True,
        key=f"close_preview_{preview_state_key}",
    ):
        st.session_state[preview_state_key] = None
        st.session_state[table_version_key] = (
            int(st.session_state.get(table_version_key, 0)) + 1
        )
        st.rerun()
