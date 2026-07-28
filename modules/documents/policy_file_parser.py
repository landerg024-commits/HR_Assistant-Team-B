"""Safe text extraction for uploaded HR policy files.

Supported formats:
- PDF (.pdf)
- Microsoft Word (.docx)
- Plain text (.txt)
- Markdown (.md)

OCR is intentionally not included yet. A scanned image-only PDF is rejected
instead of pretending that readable policy text was found.
"""

from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from pathlib import Path
import mimetypes
import re

from docx import Document
from pypdf import PdfReader


ALLOWED_POLICY_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}


@dataclass(slots=True)
class ParsedPolicySection:
    """One ordered section extracted from a source file."""

    heading: str
    text: str
    page_number: int | None = None


@dataclass(slots=True)
class ParsedPolicyDocument:
    """Validated file metadata and extracted searchable text."""

    original_filename: str
    file_extension: str
    mime_type: str
    sha256: str
    size_bytes: int
    page_count: int | None
    full_text: str
    sections: list[ParsedPolicySection]


class PolicyFileParser:
    """Validate an uploaded policy file and extract readable sections."""

    @staticmethod
    def _safe_original_name(filename: str) -> str:
        """Return a basename without path traversal components."""

        cleaned = Path(filename or "").name.strip()

        if not cleaned:
            raise ValueError("The uploaded file has no valid filename.")

        return cleaned[:255]

    @staticmethod
    def _normalize_text(value: str) -> str:
        """Normalize line endings and repeated horizontal whitespace."""

        value = value.replace("\r\n", "\n").replace("\r", "\n")
        value = re.sub(r"[ \t]+", " ", value)
        value = re.sub(r"\n{3,}", "\n\n", value)

        return value.strip()

    @staticmethod
    def _looks_like_heading(line: str) -> bool:
        """Detect generic headings without domain-specific hardcoding."""

        stripped = line.strip()

        if not stripped or len(stripped) > 140:
            return False

        if stripped.startswith("#"):
            return True

        if stripped.endswith(":"):
            return True

        if re.match(
            r"^(section|article|part|chapter)\s+[\w.-]+",
            stripped,
            flags=re.IGNORECASE,
        ):
            return True

        letters = [
            character
            for character in stripped
            if character.isalpha()
        ]

        if letters:
            uppercase_ratio = sum(
                character.isupper()
                for character in letters
            ) / len(letters)

            if uppercase_ratio >= 0.85 and len(stripped) >= 4:
                return True

        if re.match(
            r"^\d+(?:\.\d+)*[.)]?\s+\S+",
            stripped,
        ):
            return True

        return False

    @classmethod
    def _sections_from_lines(
        cls,
        lines: list[str],
        *,
        default_heading: str,
        page_number: int | None,
    ) -> list[ParsedPolicySection]:
        """Build readable sections and cap oversized section text."""

        sections: list[ParsedPolicySection] = []
        heading = default_heading
        body_lines: list[str] = []

        def flush() -> None:
            nonlocal body_lines

            text = cls._normalize_text(
                "\n".join(body_lines)
            )

            if text:
                # Split long sections by paragraph to keep retrieval focused.
                paragraphs = [
                    paragraph.strip()
                    for paragraph in re.split(
                        r"\n\s*\n",
                        text,
                    )
                    if paragraph.strip()
                ]

                chunk = ""
                part_number = 1

                for paragraph in paragraphs or [text]:
                    candidate = (
                        f"{chunk}\n\n{paragraph}".strip()
                    )

                    if len(candidate) > 1800 and chunk:
                        display_heading = (
                            heading
                            if part_number == 1
                            else f"{heading} — Part {part_number}"
                        )
                        sections.append(
                            ParsedPolicySection(
                                heading=display_heading[:250],
                                text=chunk,
                                page_number=page_number,
                            )
                        )
                        part_number += 1
                        chunk = paragraph
                    else:
                        chunk = candidate

                if chunk:
                    display_heading = (
                        heading
                        if part_number == 1
                        else f"{heading} — Part {part_number}"
                    )
                    sections.append(
                        ParsedPolicySection(
                            heading=display_heading[:250],
                            text=chunk,
                            page_number=page_number,
                        )
                    )

            body_lines = []

        for raw_line in lines:
            line = raw_line.strip()

            if not line:
                body_lines.append("")
                continue

            if cls._looks_like_heading(line):
                flush()
                heading = (
                    line.lstrip("#")
                    .strip()
                    .rstrip(":")
                    or default_heading
                )
                continue

            body_lines.append(line)

        flush()

        return sections

    @classmethod
    def parse_edited_content(
        cls,
        *,
        content: str,
        default_heading: str,
    ) -> tuple[
        str,
        list[ParsedPolicySection],
    ]:
        """Normalize edited text and regenerate searchable sections.

        The original uploaded source file is not changed. These returned
        values are used only for database content, the administrator viewer,
        and Policy Q&A.
        """

        normalized = cls._normalize_text(content)

        if len(normalized) < 20:
            raise ValueError(
                "Policy content must contain at least 20 readable "
                "characters."
            )

        normalized_heading = (
            cls._normalize_text(default_heading)
            or "Policy Details"
        )

        sections = cls._sections_from_lines(
            normalized.splitlines(),
            default_heading=normalized_heading[:250],
            page_number=None,
        )

        if not sections:
            sections = [
                ParsedPolicySection(
                    heading=normalized_heading[:250],
                    text=normalized,
                    page_number=None,
                )
            ]

        return normalized, sections

    @classmethod
    def _parse_pdf(
        cls,
        file_bytes: bytes,
    ) -> tuple[str, list[ParsedPolicySection], int]:
        """Extract page-aware sections from a text-based PDF."""

        reader = PdfReader(BytesIO(file_bytes))

        if reader.is_encrypted:
            try:
                result = reader.decrypt("")
            except Exception as error:
                raise ValueError(
                    "Encrypted PDF files are not supported."
                ) from error

            if result == 0:
                raise ValueError(
                    "Encrypted PDF files are not supported."
                )

        full_pages: list[str] = []
        sections: list[ParsedPolicySection] = []

        for page_index, page in enumerate(
            reader.pages,
            start=1,
        ):
            extracted = cls._normalize_text(
                page.extract_text() or ""
            )

            if not extracted:
                continue

            full_pages.append(
                f"[Page {page_index}]\n{extracted}"
            )
            sections.extend(
                cls._sections_from_lines(
                    extracted.splitlines(),
                    default_heading=f"Page {page_index}",
                    page_number=page_index,
                )
            )

        if not full_pages:
            raise ValueError(
                "No readable text was found in the PDF. "
                "Image-only or scanned PDFs require OCR, "
                "which is not enabled yet."
            )

        return (
            "\n\n".join(full_pages),
            sections,
            len(reader.pages),
        )

    @classmethod
    def _parse_docx(
        cls,
        file_bytes: bytes,
    ) -> tuple[str, list[ParsedPolicySection], None]:
        """Extract headings, paragraphs, and table rows from DOCX."""

        try:
            document = Document(BytesIO(file_bytes))
        except Exception as error:
            raise ValueError(
                "The DOCX file is damaged or unreadable."
            ) from error

        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if not text:
                lines.append("")
                continue

            style_name = (
                paragraph.style.name.lower()
                if paragraph.style
                and paragraph.style.name
                else ""
            )

            if style_name.startswith("heading"):
                lines.append(f"# {text}")
            else:
                lines.append(text)

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            lines.append(f"# Table {table_index}")

            for row in table.rows:
                values = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                row_text = " | ".join(
                    value
                    for value in values
                    if value
                )

                if row_text:
                    lines.append(row_text)

        full_text = cls._normalize_text(
            "\n".join(lines)
        )

        if not full_text:
            raise ValueError(
                "No readable text was found in the DOCX file."
            )

        sections = cls._sections_from_lines(
            lines,
            default_heading="Policy Details",
            page_number=None,
        )

        return full_text, sections, None

    @classmethod
    def _parse_text(
        cls,
        file_bytes: bytes,
    ) -> tuple[str, list[ParsedPolicySection], None]:
        """Decode UTF-8 text, with a Windows text fallback."""

        decoded = None

        for encoding in (
            "utf-8-sig",
            "utf-8",
            "cp1252",
        ):
            try:
                decoded = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if decoded is None:
            raise ValueError(
                "The text file encoding is not supported."
            )

        full_text = cls._normalize_text(decoded)

        if not full_text:
            raise ValueError(
                "The uploaded text file is empty."
            )

        sections = cls._sections_from_lines(
            full_text.splitlines(),
            default_heading="Policy Details",
            page_number=None,
        )

        return full_text, sections, None

    @classmethod
    def parse(
        cls,
        *,
        filename: str,
        file_bytes: bytes,
        maximum_size_bytes: int,
        supplied_mime_type: str | None = None,
    ) -> ParsedPolicyDocument:
        """Validate size/type and extract searchable policy text."""

        original_filename = cls._safe_original_name(filename)
        extension = Path(original_filename).suffix.lower()

        if extension not in ALLOWED_POLICY_EXTENSIONS:
            raise ValueError(
                "Unsupported policy file type. "
                "Upload PDF, DOCX, TXT, or MD."
            )

        if not file_bytes:
            raise ValueError("The uploaded policy file is empty.")

        if len(file_bytes) > maximum_size_bytes:
            maximum_mb = maximum_size_bytes // (1024 * 1024)
            raise ValueError(
                f"The uploaded file exceeds the {maximum_mb} MB limit."
            )

        if extension == ".pdf":
            full_text, sections, page_count = cls._parse_pdf(
                file_bytes
            )
        elif extension == ".docx":
            full_text, sections, page_count = cls._parse_docx(
                file_bytes
            )
        else:
            full_text, sections, page_count = cls._parse_text(
                file_bytes
            )

        if not sections:
            sections = [
                ParsedPolicySection(
                    heading="Policy Details",
                    text=full_text,
                    page_number=None,
                )
            ]

        guessed_mime_type = (
            mimetypes.guess_type(original_filename)[0]
            or "application/octet-stream"
        )

        return ParsedPolicyDocument(
            original_filename=original_filename,
            file_extension=extension,
            mime_type=(
                supplied_mime_type
                or guessed_mime_type
            )[:150],
            sha256=sha256(file_bytes).hexdigest(),
            size_bytes=len(file_bytes),
            page_count=page_count,
            full_text=full_text,
            sections=sections,
        )
