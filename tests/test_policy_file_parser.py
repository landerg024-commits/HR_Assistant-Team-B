"""Tests for policy file validation and extraction."""

from io import BytesIO

from docx import Document
from pypdf import PdfWriter

from modules.documents.policy_file_parser import (
    PolicyFileParser,
)


MAX_BYTES = 2 * 1024 * 1024


def test_txt_policy_extracts_sections() -> None:
    parsed = PolicyFileParser.parse(
        filename="Leave Policy.txt",
        file_bytes=(
            b"ENTITLEMENT:\n"
            b"Employees receive fifteen leave days.\n\n"
            b"REQUESTS:\n"
            b"Submit five days in advance."
        ),
        maximum_size_bytes=MAX_BYTES,
        supplied_mime_type="text/plain",
    )

    assert parsed.original_filename == "Leave Policy.txt"
    assert parsed.file_extension == ".txt"
    assert len(parsed.sections) == 2
    assert parsed.sections[0].heading == "ENTITLEMENT"


def test_docx_policy_preserves_heading() -> None:
    document = Document()
    document.add_heading("Annual Leave", level=1)
    document.add_paragraph(
        "Employees receive fifteen annual leave days."
    )

    buffer = BytesIO()
    document.save(buffer)

    parsed = PolicyFileParser.parse(
        filename="leave.docx",
        file_bytes=buffer.getvalue(),
        maximum_size_bytes=MAX_BYTES,
    )

    assert parsed.sections[0].heading == "Annual Leave"
    assert "fifteen annual leave days" in parsed.full_text


def test_blank_pdf_is_rejected_without_ocr() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)

    try:
        PolicyFileParser.parse(
            filename="scan.pdf",
            file_bytes=buffer.getvalue(),
            maximum_size_bytes=MAX_BYTES,
        )
    except ValueError as error:
        assert "No readable text" in str(error)
    else:
        raise AssertionError("Blank PDF was accepted.")


def test_unsupported_extension_is_rejected() -> None:
    try:
        PolicyFileParser.parse(
            filename="policy.exe",
            file_bytes=b"not allowed",
            maximum_size_bytes=MAX_BYTES,
        )
    except ValueError as error:
        assert "Unsupported policy file type" in str(error)
    else:
        raise AssertionError("Unsupported file was accepted.")


def test_oversized_file_is_rejected() -> None:
    try:
        PolicyFileParser.parse(
            filename="large.txt",
            file_bytes=b"x" * 100,
            maximum_size_bytes=50,
        )
    except ValueError as error:
        assert "exceeds" in str(error)
    else:
        raise AssertionError("Oversized file was accepted.")
